import os
from typing import Dict


def extract_pdf_text(file_path: str) -> str:
    text_parts = []

    try:
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text_parts.append(page.extract_text() or "")
    except Exception:
        text_parts = []

    if text_parts and any(part.strip() for part in text_parts):
        return "\n".join(text_parts).strip()

    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(file_path)
        for page in reader.pages:
            text_parts.append(page.extract_text() or "")
    except Exception:
        return ""

    return "\n".join(text_parts).strip()


def extract_docx_text(file_path: str) -> str:
    try:
        from docx import Document
    except Exception:
        return ""

    document = Document(file_path)
    text_parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]

    for table in document.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(" | ".join(row_text))

    return "\n".join(text_parts).strip()


def extract_resume_text(file_path: str) -> Dict[str, str]:
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        resume_text = extract_pdf_text(file_path)
    elif ext == ".docx":
        resume_text = extract_docx_text(file_path)
    else:
        resume_text = ""

    return {"resume_text": resume_text}
